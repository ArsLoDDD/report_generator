from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path(__file__).resolve().parents[1] / "src-tauri" / "templates"


def format_run(run, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def add_paragraph(document, text="", alignment=None, before=0, after=6, bold=False, size=12):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    if alignment is not None:
        paragraph.alignment = alignment
    format_run(paragraph.add_run(text), bold, size)
    return paragraph


def base_document(title):
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(12)
    add_paragraph(document, "КОМАНДИРУ ВІЙСЬКОВОЇ ЧАСТИНИ", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_paragraph(document, "від {{військовий_1_звання:родовий}} {{військовий_1_піб:родовий}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_paragraph(document, "{{військовий_1_посада:родовий}}", WD_ALIGN_PARAGRAPH.RIGHT, after=14)
    add_paragraph(document, "РАПОРТ", WD_ALIGN_PARAGRAPH.CENTER, after=5, bold=True, size=14)
    add_paragraph(document, title, WD_ALIGN_PARAGRAPH.CENTER, after=14, size=12)
    return document


def save_vacation(with_date=False):
    document = base_document("про надання відпустки")
    add_paragraph(document, "Прошу надати мені щорічну основну відпустку відповідно до законодавства України.", after=10)
    if with_date:
        add_paragraph(document, "Дата рапорту: {{дата_рапорту}}", after=4)
    add_paragraph(document, "Дата народження: {{військовий_1_дата_народження}}", after=4)
    add_paragraph(document, "ІПН: {{військовий_1_іпн}}", after=20)
    add_paragraph(document, "{{основний_підписант_звання}} {{основний_підписант_піб}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_paragraph(document, "{{основний_підписант_посада}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    document.save(OUTPUT / ("Рапорт на відпустку з датою.docx" if with_date else "Рапорт на відпустку.docx"))


def save_material_assistance():
    document = base_document("про надання матеріальної допомоги")
    add_paragraph(document, "Прошу розглянути питання щодо надання мені матеріальної допомоги.", after=10)
    add_paragraph(document, "Освіта: {{військовий_1_освіта}}. {{військовий_1_де_отримана_освіта}}", after=4)
    add_paragraph(document, "Служба в ЗСУ: {{військовий_1_служба_в_зсу}}", after=4)
    add_paragraph(document, "Призначений на посаду: {{військовий_1_дата_призначення}}", after=4)
    add_paragraph(document, "Наказ: {{військовий_1_наказ_призначення}}", after=20)
    add_paragraph(document, "{{основний_підписант_звання}} {{основний_підписант_піб}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_paragraph(document, "{{основний_підписант_посада}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    document.save(OUTPUT / "Рапорт на матеріальну допомогу.docx")


def save_personnel_list():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    add_paragraph(document, "СПИСОК ВІЙСЬКОВОСЛУЖБОВЦІВ", WD_ALIGN_PARAGRAPH.CENTER, after=14, bold=True, size=14)
    add_paragraph(document, "1. {{військовий_1_звання}} {{військовий_1_піб}}, {{військовий_1_посада}}", after=4)
    add_paragraph(document, "2. {{військовий_2_звання}} {{військовий_2_піб}}, {{військовий_2_посада}}", after=4)
    add_paragraph(document, "3. {{військовий_3_звання}} {{військовий_3_піб}}, {{військовий_3_посада}}", after=16)
    add_paragraph(document, "Командир: {{командир_піб}}", WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    document.save(OUTPUT / "Список військовослужбовців.docx")


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    save_vacation()
    save_vacation(with_date=True)
    save_material_assistance()
    save_personnel_list()


if __name__ == "__main__":
    main()
